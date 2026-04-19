import os
from dotenv import load_dotenv
from docx import Document
from langgraph.types import Command
from pypdf import PdfReader

from workflow import create_workflow
from graph_state import GraphState

load_dotenv()

graph = create_workflow()
config = {"configurable": {"thread_id": "1"}}


def _load_docx(path: str) -> str:
    doc = Document(path)
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def load_document(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        reader = PdfReader(path)
        return "".join(page.extract_text() or "" for page in reader.pages)
    if ext == ".docx":
        return _load_docx(path)
    raise ValueError(
        f"Unsupported format: {ext or '(no extension)'} — use .pdf or .docx"
    )


def run():

    USE_REFERENCE_FILES = False # True to use reference files, False to use zero-data mode
    ref_paths = (
        [
            r"C:\Users\hosam\OneDrive\سطح المكتب\قياس\a.docx",
            r"C:\Users\hosam\OneDrive\سطح المكتب\قياس\b.docx",
            r"C:\Users\hosam\OneDrive\سطح المكتب\قياس\c.docx",
        ]
        if USE_REFERENCE_FILES
        else []
    )
    print("Loading reference files...")

    parts = []

    for p in ref_paths:
        if not os.path.exists(p):
            print(f"Reference file not found: {p}")
            continue
        
        file_text = load_document(p)

        section = (
            f"[SOURCE: {os.path.basename(p)}]\n"
            f"{file_text}\n"
            f"[END SOURCE: {os.path.basename(p)}]"
        )
        parts.append(section)
    
    has_references = len(parts) > 0

    if has_references:
        references_content = "\n\n" + ("-" * 80) + "\n\n"
        references_content = references_content.join(parts)

        print(f"Read {len(references_content)} characters from {len(parts)} reference files.")
    else:

        references_content = ""
        print("No references uploaded. Switching to zero-data mode.")

    template_path = r"C:\Users\hosam\OneDrive\سطح المكتب\تمبلت.docx"
    
    general_project_context = "the project was in 2000 and the project name is وكيل"



    # Detect mode based on whether template file exists
    if  has_references and os.path.exists(template_path):
        print("Template found — using template mode.")
        template_text = load_document(template_path)
        mode = "template"
    elif has_references:

        print("No template, but references exist — using manual mode.")
        template_text = ""
        mode = "manual"
    
    else:
        print("No references or template — using zero-data mode.")
        template_text = ""
        mode = "zero_data"

    initial_state = GraphState(
        references_content=references_content,
        chapter_samples="",
        template=template_text,
        general_project_context=general_project_context,
        mode=mode,
        doc_type="",
        requested_chapter_count=0,
        topic="",
        chapters=[],
        current_chapter_index=0,
        draft="",
        feedback=[],
        action="",
        generated_chapters=[],
    )

    result = graph.invoke(initial_state, config=config)

    while True:
        state = graph.get_state(config)
        # Pending human-in-the-loop prompts live on state.interrupts (aggregated).
        # After a resume, next can be empty while the same node still has another
        # interrupt (e.g. collect_input asks topic, then type, then count) — do not
        # treat that as "finished".
        interrupts = tuple(state.interrupts) if state.interrupts else ()

        if interrupts:
            prompt_text = interrupts[0].value
            print("\n" + "=" * 60)
            print(prompt_text)
            print("=" * 60)

            user_input = input("\nYour reply: ").strip()
            graph.invoke(
                Command(resume=user_input),
                config=config,
            )
            continue

        if not state.next:
            print("\nContent generation finished!")
            break

        break


if __name__ == "__main__":
    run()
